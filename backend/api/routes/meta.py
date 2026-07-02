"""跨项目 / 跨 run 的元数据接口。

涵盖：
- 全局指标聚合（``/api/metrics/aggregate``）
- 单项目指标时间序列（``/api/projects/{id}/metrics/timeseries``）
- LLM call 流水（``/api/projects/{id}/llm-calls`` + ``/api/llm-calls``）
- 导出（``/api/projects/{id}/export?format=json|markdown``）
"""

from __future__ import annotations

import json
from datetime import UTC, datetime
from typing import Any, Literal

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import PlainTextResponse, Response
from pydantic import BaseModel, ConfigDict

from backend.api.deps import (
    get_current_user,
    get_owned_project,
    get_storage,
)
from backend.observability.llm_call_log import list_calls
from backend.orchestrator.metrics import best_round_reporter_key
from backend.schemas import Project, ReporterOutput, User
from backend.schemas.labels import industry_label
from backend.schemas.project import ProjectMetricsSnapshot
from backend.storage import Storage

router = APIRouter(tags=["meta"])


# ============================================================
# 全局指标聚合
# ============================================================


class AggregateMetricsResponse(BaseModel):
    model_config = ConfigDict(extra="forbid")

    project_count: int
    finished_project_count: int

    avg_accuracy: float
    avg_coverage: float
    avg_edit_rate: float
    total_evidence: int
    total_tokens: int
    total_cost_usd: float
    total_duration_seconds: int
    total_qa_rounds: int
    total_manual_edits: int

    by_status: dict[str, int]
    by_industry: dict[str, int]


@router.get("/metrics/aggregate", response_model=AggregateMetricsResponse)
async def aggregate_metrics(
    storage: Storage = Depends(get_storage),
    current_user: User = Depends(get_current_user),
    since_iso: str | None = Query(
        default=None, description="ISO8601；只算 created_at ≥ 此值的项目"
    ),
) -> AggregateMetricsResponse:
    """跨项目汇总（仅当前用户自己的项目）。"""
    projects = await storage.state_store.list_projects(owner=current_user.user_id, limit=10000)

    since_dt: datetime | None = None
    if since_iso:
        try:
            since_dt = datetime.fromisoformat(since_iso)
        except ValueError as e:
            raise HTTPException(status_code=400, detail=f"bad since_iso: {e}") from e

    if since_dt:
        projects = [p for p in projects if p.created_at >= since_dt]

    by_status: dict[str, int] = {}
    by_industry: dict[str, int] = {}
    for p in projects:
        by_status[p.status.value] = by_status.get(p.status.value, 0) + 1
        by_industry[p.industry] = by_industry.get(p.industry, 0) + 1

    with_metrics = [p for p in projects if p.metrics is not None]
    n = len(with_metrics)
    if n == 0:
        return AggregateMetricsResponse(
            project_count=len(projects),
            finished_project_count=0,
            avg_accuracy=0.0,
            avg_coverage=0.0,
            avg_edit_rate=0.0,
            total_evidence=0,
            total_tokens=0,
            total_cost_usd=0.0,
            total_duration_seconds=0,
            total_qa_rounds=0,
            total_manual_edits=0,
            by_status=by_status,
            by_industry=by_industry,
        )

    return AggregateMetricsResponse(
        project_count=len(projects),
        finished_project_count=n,
        avg_accuracy=sum(p.metrics.accuracy for p in with_metrics) / n,
        avg_coverage=sum(p.metrics.coverage for p in with_metrics) / n,
        avg_edit_rate=sum(p.metrics.edit_rate for p in with_metrics) / n,
        total_evidence=sum(p.metrics.evidence_count for p in with_metrics),
        total_tokens=sum(p.metrics.total_tokens for p in with_metrics),
        total_cost_usd=sum(p.metrics.total_cost_usd for p in with_metrics),
        total_duration_seconds=sum(p.metrics.duration_seconds for p in with_metrics),
        total_qa_rounds=sum(p.metrics.qa_round_count for p in with_metrics),
        total_manual_edits=sum(p.metrics.manual_edits for p in with_metrics),
        by_status=by_status,
        by_industry=by_industry,
    )


# ============================================================
# 单项目时间序列
# ============================================================


class MetricsTimeseriesResponse(BaseModel):
    model_config = ConfigDict(extra="forbid")

    project_id: str
    history: list[ProjectMetricsSnapshot]


@router.get(
    "/projects/{project_id}/metrics/timeseries",
    response_model=MetricsTimeseriesResponse,
)
async def metrics_timeseries(
    project_id: str,
    storage: Storage = Depends(get_storage),
    project: Project = Depends(get_owned_project),
) -> MetricsTimeseriesResponse:
    return MetricsTimeseriesResponse(project_id=project_id, history=list(project.metrics_history))


# ============================================================
# LLM call 流水
# ============================================================


class LLMCallsResponse(BaseModel):
    model_config = ConfigDict(extra="forbid")

    calls: list[dict[str, Any]]
    total: int


@router.get(
    "/projects/{project_id}/llm-calls",
    response_model=LLMCallsResponse,
)
async def project_llm_calls(
    project_id: str,
    node_id: str | None = Query(default=None),
    agent_name: str | None = Query(default=None),
    limit: int = Query(default=200, ge=1, le=2000),
    storage: Storage = Depends(get_storage),
    _project: Project = Depends(get_owned_project),
) -> LLMCallsResponse:
    """单项目的 LLM 调用流水（按节点 / Agent 过滤）。

    数据源是进程内 ring buffer（v1 单进程演示用；多进程 / 重启场景换 OTLP / 存储后端）。
    """
    trace_id = f"trace_{project_id}"
    # 实时：本进程 ring buffer（含进行中节点，可能还没落库）
    live = [
        r.to_dict()
        for r in list_calls(trace_id=trace_id, node_id=node_id, agent_name=agent_name, limit=limit)
    ]
    # 持久化：postgres 模式重启后仍可查（memory 模式为进程内）
    persisted = await storage.state_store.list_llm_calls(
        project_id, node_id=node_id, agent_name=agent_name, limit=limit
    )
    merged = _merge_llm_calls(persisted, live, limit)
    return LLMCallsResponse(calls=merged, total=len(merged))


def _merge_llm_calls(
    persisted: list[dict[str, Any]], live: list[dict[str, Any]], limit: int
) -> list[dict[str, Any]]:
    """合并持久化 + ring buffer 两路，按 (timestamp, node_id, phase) 去重，
    timestamp 倒序，截到 limit。重启后 live 为空 → 纯持久化；运行时两路并存
    → live 补上还没落库的进行中调用。"""
    seen: set[tuple] = set()
    out: list[dict[str, Any]] = []
    for c in [*persisted, *live]:
        key = (c.get("timestamp"), c.get("node_id"), c.get("phase"))
        if key in seen:
            continue
        seen.add(key)
        out.append(c)
    out.sort(key=lambda c: c.get("timestamp") or 0.0, reverse=True)
    return out[:limit]


@router.get("/llm-calls", response_model=LLMCallsResponse)
async def all_llm_calls(
    limit: int = Query(default=500, ge=1, le=5000),
    storage: Storage = Depends(get_storage),
    current_user: User = Depends(get_current_user),
) -> LLMCallsResponse:
    """当前用户名下所有项目的 LLM 调用流水（/metrics 仪表盘看实时 token 流速）。

    ring buffer 记录用 trace_id 关联项目（trace_{project_id}）；这里按本人项目集合过滤，
    防止跨用户看到他人 token 流水。
    """
    projects = await storage.state_store.list_projects(owner=current_user.user_id, limit=10000)
    allowed_traces = {f"trace_{p.project_id}" for p in projects}
    recs = [r for r in list_calls(limit=limit) if r.trace_id in allowed_traces]
    return LLMCallsResponse(calls=[r.to_dict() for r in recs], total=len(recs))


# ============================================================
# 导出
# ============================================================


@router.get("/projects/{project_id}/export")
async def export_project(
    project_id: str,
    fmt: Literal["json", "markdown", "pdf", "docx"] = Query(default="json", alias="format"),
    include_audit: bool = Query(
        default=False,
        description="markdown/pdf/docx 末尾附加「方法论与数据可信度」附录"
        "（运行指标 + 全量 evidence 索引 + QA verdict）。默认 False，输出干净交付物。",
    ),
    storage: Storage = Depends(get_storage),
    project: Project = Depends(get_owned_project),
) -> Response:
    """导出项目最终报告 + 完整数据。

    - ``format=json``：完整 state（含 plan / outputs / verdicts / metrics），方便给到老板 / 客户做底层数据
    - ``format=markdown``：读者向报告 Markdown —— 干净正文 + 编号引用 ``[n]`` + 「数据来源」列表 + 合规声明
    - ``format=pdf``：基于 markdown 渲染成 PDF（需要安装 ``reportlab``）
    - ``format=docx``：基于 markdown 渲染成 Word docx（需要安装 ``python-docx``）

    内部 ID / QA 自评分 / token 指标默认**不进交付物**；需要溯源审计时传
    ``include_audit=true`` 在末尾追加附录。

    PDF / DOCX 依赖在可选 extras 里：``pip install '.[export-pdf-docx]'``。
    缺依赖时返回 503 而非 500，方便前端做兜底。
    """
    plan = await storage.state_store.get_dag_plan(project_id)
    outputs = await storage.state_store.list_node_outputs(project_id)
    # list_qa_verdicts 按 created_at **DESC**(最新在前)返回；而 best_round_reporter_key /
    # 审计附录 verdicts[-1] 都按「轮次升序」理解 → 这里翻成升序(round1..roundN)再用，
    # 否则会选错 reporter 版本、附录也会把最旧 verdict 当最新(P1P2-VERDICT-ORDER)。
    verdicts = list(reversed(await storage.state_store.list_qa_verdicts(project_id)))

    if fmt == "json":
        payload = {
            "project": project.model_dump(mode="json"),
            "plan": plan.model_dump(mode="json") if plan else None,
            "outputs": {nid: out.model_dump(mode="json") for nid, out in outputs.items()},
            "verdicts": [v.model_dump(mode="json") for v in verdicts],
            "exported_at": datetime.now(UTC).isoformat(),
        }
        body = json.dumps(payload, ensure_ascii=False, indent=2)
        return Response(
            content=body,
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="{project_id}.json"'},
        )

    # markdown 渲染（PDF / DOCX 都基于此做格式转换）
    md = _render_markdown(project_id, project, outputs, verdicts, include_audit=include_audit)

    if fmt == "markdown":
        return PlainTextResponse(
            content=md,
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{project_id}.md"'},
        )

    if fmt == "pdf":
        try:
            body = _render_pdf(project.project_name, md)
        except ImportError as e:
            raise HTTPException(
                status_code=503,
                detail=("PDF export requires reportlab: `pip install '.[export-pdf-docx]'`"),
            ) from e
        return Response(
            content=body,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{project_id}.pdf"'},
        )

    # docx
    try:
        body = _render_docx(project.project_name, md)
    except ImportError as e:
        raise HTTPException(
            status_code=503,
            detail=("DOCX export requires python-docx: `pip install '.[export-pdf-docx]'`"),
        ) from e
    return Response(
        content=body,
        media_type=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        headers={"Content-Disposition": f'attachment; filename="{project_id}.docx"'},
    )


# ----- markdown 渲染 -----


# evidence.source_type → 人类可读来源标签（交付物里不出现机器枚举值）
_SOURCE_TYPE_LABELS = {
    "homepage": "官网首页",
    "features": "功能页",
    "pricing_page": "官方定价页",
    "pricing": "官方定价页",
    "help_docs": "帮助文档",
    "docs": "帮助文档",
    "changelog": "更新日志",
    "customer_cases": "客户案例",
    "cases": "客户案例",
    "blog": "官方博客",
    "user_reviews": "用户评价",
    "review": "用户评价",
    "reviews": "用户评价",
    "app_market": "应用市场",
}


def _source_label(source_type: str) -> str:
    return _SOURCE_TYPE_LABELS.get(source_type, source_type or "公开来源")


def _render_markdown(project_id, project, outputs, verdicts, *, include_audit: bool = False) -> str:
    """把最终 reporter 输出渲染成「读者向」Markdown 报告。

    默认输出干净交付物：正文 + 编号引用 ``[n]`` + 「数据来源」列表 + 合规声明，
    不含任何内部 ID / QA 自评分 / token 指标。``include_audit=True`` 时在末尾追加
    「方法论与数据可信度」附录（全量 evidence 索引 + QA verdict + 运行指标），
    供数据团队 / 内部复核使用。
    """
    # 发布择优(A4)：导出维度均分最高那一轮，而非简单的最后一轮（避免返工把质量
    # 改差却照发）。无 verdict 时 best_round_reporter_key 退回「最高 revision」旧行为。
    final_key = best_round_reporter_key(outputs, verdicts)
    reporter_out = outputs.get(final_key)

    # evidence_id -> Evidence（汇聚所有 extract.* 节点的产出）。
    # 先收敛到每产品最新轮：返工重抽后 extract.X 与 extract.X_v2 并存，否则审计附录
    # 会把 v1/v2 同一产品的证据重复列出。best_round_reporter_key 仍吃完整 outputs。
    from backend.orchestrator.run_state import latest_outputs

    ev_by_id: dict[str, Any] = {}
    for nid, out in latest_outputs(outputs).items():
        if not nid.startswith("extract."):
            continue
        for ev in getattr(out, "evidences", None) or []:
            ev_by_id.setdefault(ev.evidence_id, ev)

    # 引用编号：按正文首次出现顺序分配 [1][2]...；只给能溯源到的 evidence 编号
    cite_no: dict[str, int] = {}

    def _cite(eid: str) -> int | None:
        if eid not in ev_by_id:
            return None
        if eid not in cite_no:
            cite_no[eid] = len(cite_no) + 1
        return cite_no[eid]

    lines: list[str] = []
    lines.append(f"# {project.project_name}")
    lines.append("")
    today = datetime.now(UTC).date().isoformat()
    lines.append(f"> 竞品分析报告 · {today}")
    lines.append("")
    meta_bits = [f"**目标产品**：{project.target_product}"]
    if project.competitors:
        meta_bits.append(f"**对比竞品**：{'、'.join(project.competitors)}")
    if project.industry:
        meta_bits.append(f"**行业**：{industry_label(project.industry)}")
    lines.append("　|　".join(meta_bits))
    lines.append("")

    if isinstance(reporter_out, ReporterOutput) and reporter_out.draft:
        draft = reporter_out.draft
        if draft.summary:
            lines.append("## 摘要")
            lines.append("")
            lines.append(draft.summary)
            lines.append("")
        for section in draft.sections:
            lines.append(f"## {section.title}")
            lines.append("")
            for para in section.paragraphs:
                text = (para.text or "").strip()
                if not text:
                    continue
                markers = [
                    f"[{n}]"
                    for eid in dict.fromkeys(para.evidence_ids)
                    if (n := _cite(eid)) is not None
                ]
                if markers:
                    text = f"{text} {''.join(markers)}"
                lines.append(text)
                lines.append("")
    else:
        lines.append("_（尚未产出报告草稿）_")
        lines.append("")

    # 数据来源：只列正文真正引用到的，按引用号排序，人类可读（产品 · 类型 — URL）
    if cite_no:
        lines.append("## 数据来源")
        lines.append("")
        for eid, n in sorted(cite_no.items(), key=lambda kv: kv[1]):
            ev = ev_by_id[eid]
            disputed = "（存疑）" if getattr(ev, "disputed", False) else ""
            lines.append(
                f"{n}. {ev.product_name} · {_source_label(ev.source_type)}"
                f"{disputed} — {ev.source_url}"
            )
        lines.append("")

    # 合规声明
    lines.append("---")
    lines.append("")
    lines.append(
        "> 数据来源声明：本报告基于公开 SaaS 产品官网 / 帮助文档 / 公开评论站采集，"
        "已遵守 robots.txt 与 ToS；未含个人隐私信息或非公开内容。"
        "结论为基于公开资料的 AI 分析推断，不构成投资 / 采购建议。"
    )

    if include_audit:
        lines.append("")
        lines.extend(_render_audit_appendix(project_id, project, verdicts, ev_by_id))

    return "\n".join(lines)


def _render_audit_appendix(project_id, project, verdicts, ev_by_id) -> list[str]:
    """内部复核附录：运行指标 + 全量 evidence 索引 + QA verdict。

    默认不出现在交付物里（``include_audit=True`` 才追加）。
    """
    lines: list[str] = []
    lines.append("---")
    lines.append("")
    lines.append("## 附录：方法论与数据可信度")
    lines.append("")
    lines.append("| 项 | 值 |")
    lines.append("|---|---|")
    lines.append(f"| 项目 ID | `{project_id}` |")
    lines.append(f"| 状态 | {project.status.value} |")
    if project.metrics:
        lines.append(f"| QA accuracy | {project.metrics.accuracy:.2f} |")
        lines.append(f"| Schema 覆盖率 | {project.metrics.coverage:.2f} |")
        lines.append(f"| Evidence 数 | {project.metrics.evidence_count} |")
        lines.append(f"| Token 总量 | {project.metrics.total_tokens} |")
    lines.append("")

    lines.append("### 全量 Evidence 索引")
    lines.append("")
    for ev in ev_by_id.values():
        disputed = " ⚠️ disputed" if getattr(ev, "disputed", False) else ""
        lines.append(
            f"- `{ev.evidence_id}`{disputed} · {ev.product_name} · "
            f"[{_source_label(ev.source_type)}]({ev.source_url})"
        )
        preview = (ev.content or "").strip()[:200]
        if preview:
            lines.append(f"  > {preview}{'…' if len(ev.content) > 200 else ''}")
    lines.append("")

    if verdicts:
        last = verdicts[-1]
        lines.append("### QA Verdict")
        lines.append("")
        lines.append(f"- overall_status: **{last.overall_status.value}**")
        lines.append(f"- blocking: {last.blocking}")
        for dim, res in (last.dimension_results or {}).items():
            mark = "✓" if res.pass_ else "✗"
            lines.append(f"- {mark} {dim.value}: {res.score:.2f} — {res.notes or ''}")
        lines.append("")
    return lines


# ============================================================
# PDF / DOCX 渲染（懒加载依赖；缺依赖时让上层捕 ImportError 返 503）
# ============================================================

import re as _re  # noqa: E402


def _render_pdf(title: str, md_text: str) -> bytes:
    """用 reportlab 把 markdown 渲染成 PDF bytes。

    实现策略：不引重型 markdown→HTML→PDF 流水线（避免 pango/cairo 系统依赖），
    而是直接按行解析 markdown 的最常见结构（#/##/段落/列表/引用块），用 reportlab
    flowables 拼。中文字体使用系统 STHeiti / PingFang（macOS）/ DejaVu（Linux），
    任一可用即可；都不可用时退化到 Helvetica（中文显示为方框）。
    """
    from io import BytesIO

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    font_name = _register_cjk_font(pdfmetrics, TTFont)

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=10.5,
        leading=15,
    )
    h1_style = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=20,
        leading=26,
        spaceAfter=8,
    )
    h2_style = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=15,
        leading=20,
        spaceBefore=10,
        spaceAfter=6,
    )
    quote_style = ParagraphStyle(
        "Quote",
        parent=body_style,
        leftIndent=12,
        textColor="#666666",
        fontSize=9.5,
    )

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        title=title,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )

    story = []
    for raw_line in md_text.split("\n"):
        line = raw_line.rstrip()
        if not line:
            story.append(Spacer(1, 4))
            continue
        # 分隔线 ---
        if _re.fullmatch(r"-{3,}", line):
            story.append(Spacer(1, 6))
            continue
        # 表格行 | xx | yy | —— 简化处理：当作普通段落
        if line.startswith("# "):
            story.append(Paragraph(_md_inline_to_html(line[2:].strip()), h1_style))
        elif line.startswith("## "):
            story.append(Paragraph(_md_inline_to_html(line[3:].strip()), h2_style))
        elif line.startswith("> "):
            story.append(Paragraph(_md_inline_to_html(line[2:].strip()), quote_style))
        elif line.startswith("- "):
            story.append(Paragraph("• " + _md_inline_to_html(line[2:].strip()), body_style))
        else:
            story.append(Paragraph(_md_inline_to_html(line), body_style))

    doc.build(story)
    return buf.getvalue()


def _register_cjk_font(pdfmetrics, TTFont) -> str:  # noqa: N803
    """注册一个支持中文的字体并返回 fontName。任一系统候选可用即注册成功。"""
    candidates = [
        # macOS
        ("STHeiti", "/System/Library/Fonts/STHeiti Light.ttc"),
        ("PingFang", "/System/Library/Fonts/PingFang.ttc"),
        # Linux 常见
        ("NotoSansCJK", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        ("DejaVuSans", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    import os

    for name, path in candidates:
        if not os.path.exists(path):
            continue
        try:
            pdfmetrics.registerFont(TTFont(name, path))
            return name
        except Exception:
            continue
    return "Helvetica"  # 兜底（中文会变方框，但不报错）


def _md_inline_to_html(text: str) -> str:
    """reportlab Paragraph 支持的 mini HTML 子集：**bold** → <b>；`code` → <font>。

    顺便把 < 和 & 转义掉，避免 reportlab 把它当成 XML 出错。
    """
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = _re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = _re.sub(r"`([^`]+?)`", r'<font name="Courier" color="#a40000">\1</font>', text)
    # 简单链接 [text](url) → text（reportlab 链接处理不稳，直接降级）
    text = _re.sub(r"\[(.+?)\]\(([^)]+?)\)", r"\1", text)
    return text


def _render_docx(title: str, md_text: str) -> bytes:
    """用 python-docx 把 markdown 渲染成 docx bytes。"""
    from io import BytesIO

    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)

    for raw_line in md_text.split("\n"):
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if _re.fullmatch(r"-{3,}", line):
            doc.add_paragraph("———————————————")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:].strip())
            p.style = doc.styles["Quote"] if "Quote" in doc.styles else style
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(_strip_md_inline(line))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _strip_md_inline(text: str) -> str:
    """docx 渲染不解析行内 markdown，直接把 **、` 等去掉只保留可读文本。"""
    text = _re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = _re.sub(r"`([^`]+?)`", r"\1", text)
    text = _re.sub(r"\[(.+?)\]\(([^)]+?)\)", r"\1", text)
    return text
